"""
SafestClaw Document Reader - Extract text from various formats.

Supports:
- PDF files (via PyMuPDF/fitz)
- Word documents (via python-docx)
- Plain text files
- Markdown files
- HTML files
"""

import logging
import re
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

# Try imports
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False


@dataclass
class DocumentResult:
    """Result of document extraction."""
    path: str
    format: str
    text: str
    page_count: int
    word_count: int
    char_count: int
    title: str | None = None
    author: str | None = None
    error: str | None = None


class DocumentReader:
    """
    Extract text from various document formats.

    No AI required - uses dedicated parsing libraries.
    """

    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.rtf': 'rtf',
    }

    def __init__(self, allowed_paths: list[str] | None = None):
        self.allowed_paths = [
            Path(p).expanduser().resolve()
            for p in (allowed_paths or ["~"])
        ]
        self._check_dependencies()

    def _is_allowed(self, path: Path) -> bool:
        """Check if path is within allowed directories."""
        try:
            resolved = path.expanduser().resolve()
            for allowed in self.allowed_paths:
                if resolved == allowed or resolved.is_relative_to(allowed):
                    return True
            return False
        except (OSError, ValueError):
            return False

    def _check_dependencies(self) -> dict[str, bool]:
        """Check which format dependencies are available."""
        return {
            'pdf': HAS_PYMUPDF,
            'docx': HAS_DOCX,
            'html': HAS_BS4,
            'text': True,
            'markdown': True,
        }

    def read(self, path: str | Path) -> DocumentResult:
        """
        Read a document and extract text.

        Args:
            path: Path to the document

        Returns:
            DocumentResult with extracted text and metadata
        """
        path = Path(path)

        if not self._is_allowed(path):
            return DocumentResult(
                path=str(path),
                format="unknown",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error=f"Access denied: {path} is outside allowed directories",
            )

        if not path.exists():
            return DocumentResult(
                path=str(path),
                format="unknown",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error=f"File not found: {path}",
            )

        # Determine format
        suffix = path.suffix.lower()
        format_type = self.SUPPORTED_FORMATS.get(suffix, 'unknown')

        if format_type == 'unknown':
            return DocumentResult(
                path=str(path),
                format=format_type,
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error=f"Unsupported format: {suffix}",
            )

        # Read based on format
        try:
            if format_type == 'pdf':
                return self._read_pdf(path)
            elif format_type == 'docx':
                return self._read_docx(path)
            elif format_type == 'html':
                return self._read_html(path)
            elif format_type in ('text', 'markdown'):
                return self._read_text(path, format_type)
            else:
                return DocumentResult(
                    path=str(path),
                    format=format_type,
                    text="",
                    page_count=0,
                    word_count=0,
                    char_count=0,
                    error=f"Reader not implemented for: {format_type}",
                )
        except Exception as e:
            return DocumentResult(
                path=str(path),
                format=format_type,
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error=str(e),
            )

    def _read_pdf(self, path: Path) -> DocumentResult:
        """Read PDF file using PyMuPDF."""
        if not HAS_PYMUPDF:
            return DocumentResult(
                path=str(path),
                format="pdf",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error="PyMuPDF not installed. Run: pip install pymupdf",
            )

        doc = fitz.open(path)
        text_parts = []
        page_count = len(doc)

        # Extract metadata
        metadata = doc.metadata
        title = metadata.get('title') if metadata else None
        author = metadata.get('author') if metadata else None

        # Extract text from each page
        for page in doc:
            text_parts.append(page.get_text())

        doc.close()

        text = "\n\n".join(text_parts)
        text = self._clean_text(text)

        return DocumentResult(
            path=str(path),
            format="pdf",
            text=text,
            page_count=page_count,
            word_count=len(text.split()),
            char_count=len(text),
            title=title,
            author=author,
        )

    def _read_docx(self, path: Path) -> DocumentResult:
        """Read Word document using python-docx."""
        if not HAS_DOCX:
            return DocumentResult(
                path=str(path),
                format="docx",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error="python-docx not installed. Run: pip install python-docx",
            )

        doc = DocxDocument(path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n\n".join(text_parts)
        text = self._clean_text(text)

        # Get core properties
        title = None
        author = None
        try:
            core_props = doc.core_properties
            title = core_props.title
            author = core_props.author
        except Exception:
            pass

        return DocumentResult(
            path=str(path),
            format="docx",
            text=text,
            page_count=1,  # DOCX doesn't have page concept
            word_count=len(text.split()),
            char_count=len(text),
            title=title,
            author=author,
        )

    def _read_html(self, path: Path) -> DocumentResult:
        """Read HTML file using BeautifulSoup."""
        if not HAS_BS4:
            return DocumentResult(
                path=str(path),
                format="html",
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                error="BeautifulSoup not installed. Run: pip install beautifulsoup4",
            )

        with open(path, encoding='utf-8', errors='ignore') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Get title
        title = None
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text(strip=True)

        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()

        # Get text
        text = soup.get_text(separator='\n', strip=True)
        text = self._clean_text(text)

        return DocumentResult(
            path=str(path),
            format="html",
            text=text,
            page_count=1,
            word_count=len(text.split()),
            char_count=len(text),
            title=title,
        )

    def _read_text(self, path: Path, format_type: str) -> DocumentResult:
        """Read plain text or markdown file."""
        with open(path, encoding='utf-8', errors='ignore') as f:
            text = f.read()

        # For markdown, optionally strip formatting
        if format_type == 'markdown':
            # Remove common markdown syntax for plain text
            text_clean = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)  # Headers
            text_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text_clean)  # Bold
            text_clean = re.sub(r'\*([^*]+)\*', r'\1', text_clean)  # Italic
            text_clean = re.sub(r'`([^`]+)`', r'\1', text_clean)  # Inline code
            text_clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text_clean)  # Links
        else:
            text_clean = text

        text_clean = self._clean_text(text_clean)

        return DocumentResult(
            path=str(path),
            format=format_type,
            text=text_clean,
            page_count=1,
            word_count=len(text_clean.split()),
            char_count=len(text_clean),
        )

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\t+', ' ', text)

        # Remove null characters
        text = text.replace('\x00', '')

        return text.strip()

    def get_supported_formats(self) -> list[str]:
        """Return list of supported file extensions."""
        deps = self._check_dependencies()
        supported = []

        for ext, format_type in self.SUPPORTED_FORMATS.items():
            if deps.get(format_type, False):
                supported.append(ext)

        return supported

    def read_multiple(self, paths: list[str | Path]) -> list[DocumentResult]:
        """Read multiple documents."""
        return [self.read(p) for p in paths]
